#!/usr/bin/env python3
"""
website2docs - Crawl a website and export its content to PDF, DOCX, ODT, HTML or Markdown.

Usage (examples):
  python website2docs.py --url https://example.com --format pdf --output example.pdf
  python website2docs.py -u https://example.com -f docx
  python website2docs.py -u https://example.com -f odt --max-pages 200 --max-depth 3
  python website2docs.py -u https://example.com -f html
  python website2docs.py -u https://example.com -f md

Notes:
- The crawler follows only HTTP(S) links within the same domain as the start URL.
- Query strings and URL fragments are ignored by default for deduplication (can reduce duplicates).
- PDF output includes basic text, images, lists, and tables; characters not supported by base PDF fonts will be replaced with '?'.
- Requires Python 3.8+ (tested for 3.13 compatibility) and the dependencies listed in requirements.txt.
"""
from __future__ import annotations

import argparse
import dataclasses
import datetime as _dt
import re
import sys
import time
from collections import deque
from dataclasses import dataclass
from typing import Iterable, List, Optional, Set, Tuple

from urllib.parse import urljoin, urlparse, urlunparse

# Third-party libs (declared in requirements.txt)
# Imported lazily inside functions to allow --help without dependencies.


@dataclass
class PageContent:
    url: str
    title: str
    text: str
    html: str = ""


USER_AGENT = "website2docs/1.0 (+https://github.com/)"
DEFAULT_TIMEOUT = 15


def _canonical_netloc(netloc: str) -> str:
    netloc = netloc.strip().lower()
    if netloc.startswith("www."):
        return netloc[4:]
    return netloc


def normalize_url(base_url: str, href: str, keep_query: bool = False) -> Optional[str]:
    # When href is empty, treat base_url itself as the URL to normalize
    candidate_href = href if href else base_url
    if not candidate_href:
        return None
    # filter out non-http(s) early
    if candidate_href.startswith("javascript:") or candidate_href.startswith("mailto:") or candidate_href.startswith("tel:"):
        return None
    try:
        abs_url = urljoin(base_url, href) if href else candidate_href
        parsed = urlparse(abs_url)
        if parsed.scheme not in ("http", "https"):
            return None
        scheme = parsed.scheme.lower()
        netloc = parsed.netloc.lower()
        path = parsed.path or "/"
        # normalize duplicate slashes
        path = re.sub(r"/+", "/", path)
        # remove trailing slash (except root)
        if len(path) > 1 and path.endswith("/"):
            path = path[:-1]
        # ignore fragments
        fragment = ""
        # optionally strip query
        query = parsed.query if keep_query else ""
        normalized = urlunparse((scheme, netloc, path, "", query, fragment))
        return normalized
    except Exception:
        return None


def is_same_domain(start_url: str, other_url: str) -> bool:
    a = _canonical_netloc(urlparse(start_url).netloc)
    b = _canonical_netloc(urlparse(other_url).netloc)
    return a == b


def extract_confluence_info(url: str) -> Tuple[Optional[str], Optional[str]]:
    """Extract domain and page ID from a Confluence URL."""
    parsed = urlparse(url)
    if not parsed.netloc:
        return None, None
    domain = f"{parsed.scheme}://{parsed.netloc}"
    if not parsed.netloc.endswith(".atlassian.net"):
        return domain, None

    # Try to find pageId in query (e.g., viewpage.action?pageId=123)
    from urllib.parse import parse_qs
    qs = parse_qs(parsed.query)
    if "pageId" in qs:
        return domain, qs["pageId"][0]

    # Try to find in path: /wiki/spaces/KEY/pages/12345/Title
    match = re.search(r"/pages/(\d+)", parsed.path)
    if match:
        return domain, match.group(1)

    return domain, None


def crawl_confluence(
    start_url: str,
    email: str,
    token: str,
    max_depth: int = 2,
    max_pages: int = 100,
    delay: float = 0.0,
    timeout: int = DEFAULT_TIMEOUT,
) -> List[PageContent]:
    """Crawl Confluence using REST API v2."""
    try:
        import requests
    except Exception as e:
        raise RuntimeError("requests is required for Confluence mode. Install with: pip install requests") from e

    domain, start_page_id = extract_confluence_info(start_url)
    if not start_page_id:
        print(f"[website2docs] Could not extract Page ID from {start_url}", file=sys.stderr)
        return []

    auth = (email, token)
    q: deque[Tuple[str, int]] = deque([(start_page_id, 0)])
    seen: Set[str] = set()
    results: List[PageContent] = []

    while q and (max_pages <= 0 or len(results) < max_pages):
        page_id, depth = q.popleft()
        if page_id in seen:
            continue
        seen.add(page_id)

        # Fetch page content
        # API v2: /wiki/api/v2/pages/{id}?body-format=storage
        api_url = f"{domain}/wiki/api/v2/pages/{page_id}?body-format=storage"
        try:
            resp = requests.get(api_url, auth=auth, timeout=timeout)
            if resp.status_code != 200:
                print(f"[website2docs] Failed to fetch page {page_id}: {resp.status_code}")
                continue

            data = resp.json()
            title = data.get("title", f"Page {page_id}")
            body_storage = data.get("body", {}).get("storage", {}).get("value", "")

            # Confluence Storage format is XHTML.
            # We can use our existing extract_text_and_title to get plain text.
            _, text = extract_text_and_title(body_storage)

            # Optional: handle images by converting them to data URIs
            # to preserve auth when saving to DOCX/PDF later.
            if body_storage:
                try:
                    from bs4 import BeautifulSoup
                    import base64
                    soup = BeautifulSoup(body_storage, "lxml")
                    img_tags = soup.find_all("img")
                    for img in img_tags:
                        src = img.get("src")
                        if src:
                            # resolve relative URLs
                            if src.startswith("/"):
                                src = f"{domain}{src}"
                            try:
                                i_resp = requests.get(src, auth=auth, timeout=timeout)
                                if i_resp.status_code == 200:
                                    content_type = i_resp.headers.get("Content-Type", "image/png")
                                    encoded = base64.b64encode(i_resp.content).decode("utf-8")
                                    img["src"] = f"data:{content_type};base64,{encoded}"
                            except Exception:
                                pass
                    body_storage = str(soup)
                except Exception:
                    pass

            results.append(PageContent(
                url=f"{domain}/wiki/pages/viewpage.action?pageId={page_id}",
                title=title,
                text=text,
                html=body_storage
            ))

            if depth < max_depth:
                # Get children
                # API v2: /wiki/api/v2/pages/{id}/children
                child_url = f"{domain}/wiki/api/v2/pages/{page_id}/children"
                child_resp = requests.get(child_url, auth=auth, timeout=timeout)
                if child_resp.status_code == 200:
                    children = child_resp.json().get("results", [])
                    for child in children:
                        child_id = child.get("id")
                        if child_id and child_id not in seen:
                            q.append((child_id, depth + 1))
        except Exception as e:
            print(f"[website2docs] Error processing Confluence page {page_id}: {e}")

        if delay > 0:
            time.sleep(delay)

    return results


def extract_links(html: str, base_url: str, keep_query: bool = False) -> List[str]:
    try:
        from bs4 import BeautifulSoup
    except Exception as e:
        raise RuntimeError("beautifulsoup4 is required. Install with: pip install beautifulsoup4 lxml") from e
    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")
    links = []
    for a in soup.find_all("a", href=True):
        u = normalize_url(base_url, a.get("href", ""), keep_query=keep_query)
        if u:
            links.append(u)
    return links


def extract_text_and_title(html: str) -> Tuple[str, str]:
    """Extract a readable title and text from the HTML using simple heuristics."""
    try:
        from bs4 import BeautifulSoup
    except Exception as e:
        raise RuntimeError("beautifulsoup4 is required. Install with: pip install beautifulsoup4 lxml") from e
    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")
    # remove scripts, styles, and common non-content sections
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe", "form"]):
        tag.decompose()

    # try to focus on main content area if available
    content_root = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", id=re.compile(r"content|main", re.I))
        or soup.find("section", id=re.compile(r"content|main", re.I))
        or soup.body
        or soup
    )

    # remove navigational elements inside content_root
    for tag in content_root.find_all(["nav", "aside", "footer", "header"]):
        tag.decompose()

    # Determine title
    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    if not title:
        # fallback: first h1/h2 text
        h = content_root.find(["h1", "h2"]) if content_root else None
        if h and h.get_text(strip=True):
            title = h.get_text(strip=True)
    if not title:
        title = "Untitled"

    # Build text with line breaks between blocks
    blocks = []
    # Prefer headings and paragraphs/lists
    preferred = content_root.find_all(["h1", "h2", "h3", "h4", "p", "li"]) if content_root else []
    if preferred:
        for el in preferred:
            txt = el.get_text(" ", strip=True)
            if txt:
                if el.name in {"h1", "h2", "h3", "h4"}:
                    blocks.append(f"\n# {txt}\n")
                else:
                    blocks.append(txt)
    else:
        # fallback: all text from content_root
        txt = content_root.get_text("\n", strip=True) if content_root else soup.get_text("\n", strip=True)
        if txt:
            blocks.append(txt)

    text = "\n\n".join(b.strip() for b in blocks if b and b.strip())
    return title, text


def fetch(url: str, user_agent: str, timeout: int = DEFAULT_TIMEOUT) -> Tuple[int, str, str]:
    """Return (status_code, final_url, text_html or '')"""
    headers = {"User-Agent": user_agent, "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8"}
    try:
        import requests
    except Exception as e:
        raise RuntimeError("requests is required for crawling. Install with: pip install requests") from e
    try:
        resp = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
        ct = resp.headers.get("content-type", "").lower()
        if resp.status_code == 200 and ("text/html" in ct or "application/xhtml+xml" in ct):
            return resp.status_code, resp.url, resp.text
        return resp.status_code, resp.url, ""
    except requests.RequestException:
        return 0, url, ""


def crawl(
    start_url: str,
    max_depth: int = 2,
    max_pages: int = 100,
    delay: float = 0.0,
    user_agent: str = USER_AGENT,
    keep_query: bool = False,
    timeout: int = DEFAULT_TIMEOUT,
) -> List[PageContent]:
    start_url = normalize_url(start_url, "", keep_query=False) or start_url
    q: deque[Tuple[str, int]] = deque()
    q.append((start_url, 0))
    seen: Set[str] = set()
    results: List[PageContent] = []

    while q and (max_pages <= 0 or len(results) < max_pages):
        url, depth = q.popleft()
        # normalize the URL by itself to keep consistent (drop queries/fragments)
        url = normalize_url(url, "", keep_query=keep_query) or url
        if url in seen:
            continue
        seen.add(url)

        status, final_url, html = fetch(url, user_agent=user_agent, timeout=timeout)
        if status == 200 and html:
            title, text = extract_text_and_title(html)
            results.append(PageContent(url=final_url or url, title=title, text=text, html=html))
            # enqueue child links if depth allows
            if depth < max_depth:
                for link in extract_links(html, base_url=final_url or url, keep_query=keep_query):
                    if is_same_domain(start_url, link):
                        nurl = normalize_url(link, "", keep_query=keep_query) or link
                        if nurl not in seen:
                            q.append((nurl, depth + 1))
        # politeness delay
        if delay > 0:
            try:
                time.sleep(delay)
            except KeyboardInterrupt:
                break

    return results


# ============== Rich Content Extraction & Helpers ==============

def extract_blocks(html: str, base_url: str) -> List[dict]:
    """Parse HTML and return an ordered list of content blocks (headings, paragraphs, images, tables, lists, code).
    Each block is a dict with a 'type' field and relevant data.
    """
    try:
        from bs4 import BeautifulSoup
    except Exception as e:
        raise RuntimeError("beautifulsoup4 is required. Install with: pip install beautifulsoup4 lxml") from e
    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")

    # remove scripts, styles, and common non-content sections
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe", "form"]):
        tag.decompose()

    content_root = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", id=re.compile(r"content|main", re.I))
        or soup.find("section", id=re.compile(r"content|main", re.I))
        or soup.body
        or soup
    )

    # remove navigational elements inside content_root
    for tag in content_root.find_all(["nav", "aside", "footer", "header"]):
        tag.decompose()

    blocks: List[dict] = []

    # Collect blocks in document order
    candidates = content_root.find_all([
        "h1", "h2", "h3", "h4",
        "p",
        "img",
        "table",
        "ul", "ol",
        "pre", "code",
    ], recursive=True)

    def abs_url(u: str) -> Optional[str]:
        return normalize_url(base_url, u, keep_query=True) if u else None

    for el in candidates:
        name = el.name.lower()
        if name in {"h1", "h2", "h3", "h4"}:
            txt = el.get_text(" ", strip=True)
            if txt:
                blocks.append({"type": "heading", "level": int(name[1]), "text": txt})
        elif name == "p":
            # avoid duplicating list item paragraphs
            if el.find_parent("li") is not None or el.find_parent("table") is not None:
                continue
            txt = el.get_text(" ", strip=True)
            if txt:
                blocks.append({"type": "paragraph", "text": txt})
        elif name == "img":
            src = el.get("src")
            alt = el.get("alt", "")
            if src:
                if src.strip().lower().startswith("data:"):
                    blocks.append({"type": "image", "src": src.strip(), "alt": alt, "is_data": True})
                else:
                    u = abs_url(src)
                    if u:
                        blocks.append({"type": "image", "src": u, "alt": alt, "is_data": False})
        elif name == "table":
            rows = []
            for tr in el.find_all("tr"):
                cells = []
                for cell in tr.find_all(["th", "td"]):
                    cells.append(cell.get_text(" ", strip=True))
                if cells:
                    rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})
        elif name in ("ul", "ol"):
            items = []
            for li in el.find_all("li", recursive=False):
                t = li.get_text(" ", strip=True)
                if t:
                    items.append(t)
            if items:
                blocks.append({"type": "list", "ordered": (name == "ol"), "items": items})
        elif name == "pre":
            code_text = el.get_text("\n", strip=False)
            if code_text is not None:
                blocks.append({"type": "code", "text": code_text})
        elif name == "code":
            # Only treat standalone code, not inside <pre>
            if el.find_parent("pre") is None:
                t = el.get_text(" ", strip=True)
                if t:
                    blocks.append({"type": "code", "text": t})

    return blocks


def _guess_image_ext_from_ct(content_type: str, url: Optional[str]) -> str:
    ct = (content_type or "").lower()
    if "image/png" in ct:
        return ".png"
    if "image/jpeg" in ct or "image/jpg" in ct:
        return ".jpg"
    if "image/gif" in ct:
        return ".gif"
    if "image/webp" in ct:
        return ".webp"
    if "image/svg" in ct:
        return ".svg"
    if url:
        m = re.search(r"\.([a-z0-9]{2,5})(?:\?|$)", url.lower())
        if m:
            return "." + m.group(1)
    return ".img"


def _decode_data_url(data_url: str) -> Tuple[Optional[bytes], str]:
    """Return (bytes, ext) for data URLs like data:image/png;base64,..."""
    import base64
    try:
        header, b64 = data_url.split(",", 1)
    except ValueError:
        return None, ".img"
    # header example: data:image/png;base64
    mediatype = ""
    if ";base64" in header:
        mediatype = header[5: header.lower().find(";base64")]
        try:
            data = base64.b64decode(b64, validate=False)
        except Exception:
            data = base64.b64decode(b64 + "==")
    else:
        mediatype = header[5:]
        try:
            data = b64.encode("utf-8")
        except Exception:
            return None, ".img"
    ext = _guess_image_ext_from_ct(mediatype, None)
    return data, ext


def fetch_image_bytes(src: str, user_agent: str = USER_AGENT, timeout: int = DEFAULT_TIMEOUT) -> Tuple[Optional[bytes], str]:
    """Fetch image bytes from http(s) URL or decode data: URL. Returns (bytes or None, ext)."""
    if not src:
        return None, ".img"
    if src.strip().lower().startswith("data:"):
        return _decode_data_url(src)
    try:
        import requests
    except Exception:
        return None, ".img"
    try:
        headers = {"User-Agent": user_agent, "Accept": "image/*,*/*;q=0.8"}
        resp = requests.get(src, headers=headers, timeout=timeout, stream=True)
        if resp.status_code == 200:
            ct = resp.headers.get("content-type", "")
            if "image" in ct or src.lower().split("?")[0].endswith((".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg")):
                data = resp.content
                ext = _guess_image_ext_from_ct(ct, src)
                return data, ext
        return None, ".img"
    except Exception:
        return None, ".img"


# ============== Renderers ==============

def _split_paragraphs(text: str) -> List[str]:
    # Split on double newlines and keep reasonably sized chunks
    paras = re.split(r"\n\s*\n", text)
    cleaned: List[str] = []
    for p in paras:
        p = p.strip()
        if not p:
            continue
        cleaned.append(p)
    return cleaned


def save_as_docx(pages: List[PageContent], output_path: str, start_url: str, orientation: str = "portrait") -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.section import WD_ORIENT
    except Exception as e:
        raise RuntimeError("python-docx is required for DOCX output. Install with: pip install python-docx") from e
    import io

    doc = Document()
    # Apply page orientation
    try:
        section = doc.sections[0]
        if (orientation or "portrait").lower().startswith("land"):
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width/height for landscape
            section.page_width, section.page_height = section.page_height, section.page_width
        else:
            section.orientation = WD_ORIENT.PORTRAIT
    except Exception:
        pass
    # Title
    title = f"Website export: {start_url}"
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated on: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    for i, page in enumerate(pages, 1):
        if i > 1:
            doc.add_page_break()
        doc.add_heading(page.title or "Untitled", level=1)
        doc.add_paragraph(page.url)

        # Try rich content blocks first
        blocks: List[dict] = []
        if getattr(page, "html", ""):
            try:
                blocks = extract_blocks(page.html, base_url=page.url)
            except Exception:
                blocks = []
        if not blocks:
            for para in _split_paragraphs(page.text):
                doc.add_paragraph(para)
            continue

        for blk in blocks:
            btype = blk.get("type")
            if btype == "heading":
                lvl = int(blk.get("level", 2))
                lvl = max(1, min(4, lvl))
                doc.add_heading(blk.get("text", ""), level=lvl)
            elif btype == "paragraph":
                doc.add_paragraph(blk.get("text", ""))
            elif btype == "list":
                style = "List Number" if blk.get("ordered") else "List Bullet"
                for item in blk.get("items", []):
                    doc.add_paragraph(item, style=style)
            elif btype == "code":
                text = blk.get("text", "")
                p = doc.add_paragraph()
                lines = text.splitlines() or [text]
                for idx, line in enumerate(lines):
                    r = p.add_run(line)
                    try:
                        r.font.name = "Courier New"
                    except Exception:
                        pass
                    if idx < len(lines) - 1:
                        r.add_break()
            elif btype == "image":
                data, _ = fetch_image_bytes(blk.get("src", ""))
                if data:
                    try:
                        doc.add_picture(io.BytesIO(data), width=Inches(6.0))
                    except Exception:
                        # Fallback: add alt text if picture fails
                        alt = blk.get("alt") or "[image]"
                        doc.add_paragraph(alt)
                else:
                    alt = blk.get("alt") or "[image]"
                    doc.add_paragraph(alt)
            elif btype == "table":
                rows: List[List[str]] = blk.get("rows", [])
                if rows:
                    cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=cols)
                    for ri, row in enumerate(rows):
                        for ci in range(cols):
                            txt = row[ci] if ci < len(row) else ""
                            table.cell(ri, ci).text = txt
                    doc.add_paragraph("")  # spacing after table

    # Set a base font size (optional)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.save(output_path)


def save_as_odt(pages: List[PageContent], output_path: str, start_url: str, orientation: str = "portrait") -> None:
    try:
        from odf.opendocument import OpenDocumentText
        from odf import text as odf_text
        from odf import table as odf_table
        from odf import draw as odf_draw
        from odf import style as odf_style
    except Exception as e:
        raise RuntimeError("odfpy is required for ODT output. Install with: pip install odfpy") from e

    doc = OpenDocumentText()

    # Apply page orientation via page layout and master page (A4 default)
    try:
        is_land = (orientation or "portrait").lower().startswith("land")
        width_cm, height_cm = (29.7, 21.0) if is_land else (21.0, 29.7)
        pl = odf_style.PageLayout(name="pm_website2docs")
        plprops = odf_style.PageLayoutProperties(
            pagewidth=f"{width_cm}cm",
            pageheight=f"{height_cm}cm",
            printorientation=("landscape" if is_land else "portrait"),
        )
        pl.addElement(plprops)
        doc.automaticstyles.addElement(pl)
        mp = odf_style.MasterPage(name="Standard", pagelayoutname=pl)
        doc.masterstyles.addElement(mp)
    except Exception:
        pass

    # Temp storage for images embedded via odfpy
    import tempfile, os
    tmp_dir = tempfile.mkdtemp(prefix="website2docs_odt_")
    _tmp_files: List[str] = []

    h = odf_text.H(outlinelevel=1, text=f"Website export: {start_url}")
    doc.text.addElement(h)
    doc.text.addElement(odf_text.P(text=f"Generated on: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"))
    doc.text.addElement(odf_text.P(text=""))

    for i, page in enumerate(pages, 1):
        if i > 1:
            # ODT doesn't have an explicit page break element in odfpy text API; add a blank paragraph to separate
            doc.text.addElement(odf_text.P(text=""))
        doc.text.addElement(odf_text.H(outlinelevel=2, text=page.title or "Untitled"))
        doc.text.addElement(odf_text.P(text=page.url))

        blocks: List[dict] = []
        if getattr(page, "html", ""):
            try:
                blocks = extract_blocks(page.html, base_url=page.url)
            except Exception:
                blocks = []
        if not blocks:
            for para in _split_paragraphs(page.text):
                doc.text.addElement(odf_text.P(text=para))
            continue

        for blk in blocks:
            btype = blk.get("type")
            if btype == "heading":
                lvl = int(blk.get("level", 2))
                lvl = max(3, min(6, lvl + 1))  # keep page title at level 2; headings inside start at 3+
                doc.text.addElement(odf_text.H(outlinelevel=lvl, text=blk.get("text", "")))
            elif btype == "paragraph":
                doc.text.addElement(odf_text.P(text=blk.get("text", "")))
            elif btype == "list":
                lst = odf_text.List()
                for item in blk.get("items", []):
                    li = odf_text.ListItem()
                    li.addElement(odf_text.P(text=item))
                    lst.addElement(li)
                doc.text.addElement(lst)
            elif btype == "code":
                # Simple paragraph for code (no special style defined)
                text = blk.get("text", "")
                for line in (text.splitlines() or [text]):
                    doc.text.addElement(odf_text.P(text=line))
            elif btype == "table":
                rows: List[List[str]] = blk.get("rows", [])
                if rows:
                    t = odf_table.Table()
                    for row in rows:
                        tr = odf_table.TableRow()
                        for cell_text in row:
                            tc = odf_table.TableCell()
                            tc.addElement(odf_text.P(text=cell_text))
                            tr.addElement(tc)
                        t.addElement(tr)
                    doc.text.addElement(t)
            elif btype == "image":
                data, ext = fetch_image_bytes(blk.get("src", ""))
                if data:
                    try:
                        fname = f"img_{len(_tmp_files)+1}{ext if ext else '.img'}"
                        fpath = os.path.join(tmp_dir, fname)
                        with open(fpath, "wb") as f:
                            f.write(data)
                        _tmp_files.append(fpath)
                        href = doc.addPicture(fpath)
                        frame = odf_draw.Frame(width="16cm")
                        img = odf_draw.Image(href=href)
                        frame.addElement(img)
                        doc.text.addElement(frame)
                    except Exception:
                        alt = blk.get("alt") or "[image]"
                        doc.text.addElement(odf_text.P(text=alt))
                else:
                    alt = blk.get("alt") or "[image]"
                    doc.text.addElement(odf_text.P(text=alt))

    doc.save(output_path)


def save_as_pdf(pages: List[PageContent], output_path: str, start_url: str, orientation: str = "portrait") -> None:
    try:
        from reportlab.lib.pagesizes import A4, landscape, portrait
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, Preformatted, KeepInFrame
        from reportlab.lib.units import cm
        from reportlab.lib import fonts, colors
        from reportlab.lib.styles import ParagraphStyle
    except Exception as e:
        raise RuntimeError("reportlab is required for PDF output. Install with: pip install reportlab") from e
    import io

    # Prepare document
    pagesize = landscape(A4) if (orientation or "portrait").lower().startswith("land") else A4
    doc = SimpleDocTemplate(output_path, pagesize=pagesize, title=f"Website export: {start_url}")
    styles = getSampleStyleSheet()
    normal = styles["BodyText"]
    heading = styles["Heading1"]
    small = styles["Normal"]
    small.fontSize = 9
    code_style = ParagraphStyle('Code', parent=styles['BodyText'], fontName='Courier', fontSize=9, leading=11)
    table_cell_style = ParagraphStyle('TableCell', parent=styles['BodyText'], fontSize=9, leading=11)

    story = []

    def safe(txt: str) -> str:
        # Replace characters not representable by Latin-1 with '?'
        try:
            return txt.encode("latin-1", "replace").decode("latin-1")
        except Exception:
            return txt

    story.append(Paragraph(safe(f"Website export: {start_url}"), styles["Title"]))
    story.append(Paragraph(safe(f"Generated on: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"), small))
    story.append(Spacer(1, 0.5 * cm))

    for i, page in enumerate(pages, 1):
        if i > 1:
            story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(safe(page.title or "Untitled"), heading))
        story.append(Paragraph(safe(page.url), small))
        story.append(Spacer(1, 0.2 * cm))

        blocks: List[dict] = []
        if getattr(page, "html", ""):
            try:
                blocks = extract_blocks(page.html, base_url=page.url)
            except Exception:
                blocks = []
        if not blocks:
            for para in _split_paragraphs(page.text):
                story.append(Paragraph(safe(para).replace("\n", "<br/>"), normal))
                story.append(Spacer(1, 0.1 * cm))
            continue

        for blk in blocks:
            btype = blk.get("type")
            if btype == "heading":
                lvl = int(blk.get("level", 2))
                lvl = max(2, min(4, lvl))
                try:
                    hstyle = styles[f"Heading{lvl}"]
                except Exception:
                    hstyle = heading
                story.append(Paragraph(safe(blk.get("text", "")), hstyle))
            elif btype == "paragraph":
                story.append(Paragraph(safe(blk.get("text", "")).replace("\n", "<br/>"), normal))
            elif btype == "list":
                for item in blk.get("items", []):
                    story.append(Paragraph(safe(item), normal, bulletText=("1." if blk.get("ordered") else "\u2022")))
            elif btype == "code":
                text = blk.get("text", "")
                story.append(Preformatted(safe(text), code_style))
            elif btype == "image":
                data, _ = fetch_image_bytes(blk.get("src", ""))
                if data:
                    try:
                        img = RLImage(io.BytesIO(data))
                        # Scale image to fit within the available frame size so it can always
                        # fit on a fresh page even if it doesn't fit in the remaining space.
                        max_w, max_h = doc.width, doc.height
                        img._restrictSize(max_w, max_h)
                        # Wrap image in KeepInFrame so it shrinks to available space if needed
                        kif_img = KeepInFrame(maxWidth=doc.width, maxHeight=doc.height, content=[img], mode='shrink')
                        story.append(kif_img)
                        story.append(Spacer(1, 0.1 * cm))
                    except Exception:
                        alt = blk.get("alt") or "[image]"
                        story.append(Paragraph(safe(alt), small))
                else:
                    alt = blk.get("alt") or "[image]"
                    story.append(Paragraph(safe(alt), small))
            elif btype == "table":
                rows: List[List[str]] = blk.get("rows", [])
                if rows:
                    # Convert cell text to Paragraphs to enable wrapping
                    para_rows: List[List[Paragraph]] = []
                    max_cols = max(len(r) for r in rows)
                    for row in rows:
                        para_row: List[Paragraph] = []
                        for cell in row:
                            para_row.append(Paragraph(safe(cell), table_cell_style))
                        # Pad short rows to consistent column count
                        while len(para_row) < max_cols:
                            para_row.append(Paragraph("", table_cell_style))
                        para_rows.append(para_row)
                    # Distribute columns across available width
                    col_widths = [doc.width / max_cols for _ in range(max_cols)]
                    t = Table(para_rows, colWidths=col_widths)
                    t.setStyle(TableStyle([
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("BOX", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 2),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ]))
                    # Ensure the table never exceeds a frame; shrink if needed (e.g., very tall row)
                    kif = KeepInFrame(maxWidth=doc.width, maxHeight=doc.height, content=[t], mode='shrink')
                    story.append(kif)
                    story.append(Spacer(1, 0.1 * cm))

    doc.build(story)


def save_as_html(pages: List[PageContent], output_path: str, start_url: str, **kwargs) -> None:
    """Saves the crawled content as a single HTML file."""
    import html
    
    parts = []
    parts.append("<!DOCTYPE html>")
    parts.append("<html><head>")
    parts.append(f"<title>Website export: {html.escape(start_url)}</title>")
    parts.append("<meta charset=\"utf-8\">")
    parts.append("<style>body { font-family: sans-serif; } img { max-width: 100%; height: auto; } table { border-collapse: collapse; } td, th { border: 1px solid #ccc; padding: 4px; } </style>")
    parts.append("</head><body>")
    parts.append(f"<h1>Website export: {html.escape(start_url)}</h1>")
    parts.append(f"<p>Generated on: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")

    for page in pages:
        parts.append("<hr>")
        parts.append(f"<h2>{html.escape(page.title or 'Untitled')}</h2>")
        parts.append(f"<p><em>Source: <a href=\"{html.escape(page.url)}\">{html.escape(page.url)}</a></em></p>")
        
        blocks: List[dict] = []
        if getattr(page, "html", ""):
            try:
                blocks = extract_blocks(page.html, base_url=page.url)
            except Exception:
                blocks = []

        if not blocks:
            for para in _split_paragraphs(page.text):
                parts.append(f"<p>{html.escape(para)}</p>")
            continue

        for blk in blocks:
            btype = blk.get("type")
            if btype == "heading":
                lvl = int(blk.get("level", 2))
                parts.append(f"<h{lvl+1}>{html.escape(blk.get('text', ''))}</h{lvl+1}>")
            elif btype == "paragraph":
                parts.append(f"<p>{html.escape(blk.get('text', ''))}</p>")
            elif btype == "list":
                tag = "ol" if blk.get("ordered") else "ul"
                parts.append(f"<{tag}>")
                for item in blk.get("items", []):
                    parts.append(f"<li>{html.escape(item)}</li>")
                parts.append(f"</{tag}>")
            elif btype == "code":
                parts.append(f"<pre><code>{html.escape(blk.get('text', ''))}</code></pre>")
            elif btype == "image":
                alt_text = html.escape(blk.get('alt', ''))
                parts.append(f"<img src=\"{html.escape(blk.get('src', ''))}\" alt=\"{alt_text}\">")
            elif btype == "table":
                parts.append("<table>")
                rows: List[List[str]] = blk.get("rows", [])
                for row in rows:
                    parts.append("<tr>")
                    for cell in row:
                        parts.append(f"<td>{html.escape(cell)}</td>")
                    parts.append("</tr>")
                parts.append("</table>")

    parts.append("</body></html>")
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))


def save_as_markdown(pages: List[PageContent], output_path: str, start_url: str, **kwargs) -> None:
    """Saves the crawled content as a single Markdown file."""
    parts = []
    parts.append(f"# Website export: {start_url}")
    parts.append(f"Generated on: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for page in pages:
        parts.append("\n---\n")
        parts.append(f"## {page.title or 'Untitled'}")
        parts.append(f"*{page.url}*")
        
        blocks: List[dict] = []
        if getattr(page, "html", ""):
            try:
                blocks = extract_blocks(page.html, base_url=page.url)
            except Exception:
                blocks = []

        if not blocks:
            for para in _split_paragraphs(page.text):
                parts.append(para)
            continue

        for blk in blocks:
            btype = blk.get("type")
            if btype == "heading":
                lvl = int(blk.get("level", 2))
                parts.append(f"{'#' * (lvl + 1)} {blk.get('text', '')}")
            elif btype == "paragraph":
                parts.append(blk.get("text", ''))
            elif btype == "list":
                prefix = "1. " if blk.get("ordered") else "* "
                for item in blk.get("items", []):
                    parts.append(f"{prefix}{item}")
            elif btype == "code":
                parts.append(f"```\n{blk.get('text', '')}\n```")
            elif btype == "image":
                parts.append(f"![{blk.get('alt', '')}]({blk.get('src', '')})")
            elif btype == "table":
                rows: List[List[str]] = blk.get("rows", [])
                if rows:
                    header = "| " + " | ".join(rows[0]) + " |"
                    divider = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                    parts.append(header)
                    parts.append(divider)
                    for row in rows[1:]:
                        parts.append("| " + " | ".join(row) + " |")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(parts))


# ============== CLI ==============

def guess_format_from_extension(path: str) -> Optional[str]:
    path = path.lower()
    if path.endswith(".pdf"):
        return "pdf"
    if path.endswith(".docx"):
        return "docx"
    if path.endswith(".odt"):
        return "odt"
    if path.endswith((".html", ".htm")):
        return "html"
    if path.endswith(".md"):
        return "md"
    return None


def parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Crawl a website and export its content to PDF, DOCX, ODT, HTML, or Markdown.")
    p.add_argument("--url", "-u", required=True, help="Start URL (e.g., https://example.com)")
    p.add_argument("--format", "-f", choices=["pdf", "docx", "odt", "html", "md"], help="Output format. If omitted, inferred from --output extension.")
    p.add_argument("--output", "-o", help="Output file path. If omitted, a name is derived from the domain and format.")
    p.add_argument("--max-pages", type=int, default=100, help="Maximum number of pages to crawl (0 = unlimited). Default: 100")
    p.add_argument("--max-depth", type=int, default=2, help="Maximum crawl depth (0 = only the start page). Default: 2")
    p.add_argument("--delay", type=float, default=0.0, help="Delay between requests in seconds. Default: 0.0")
    p.add_argument("--user-agent", default=USER_AGENT, help=f"User-Agent header. Default: {USER_AGENT}")
    p.add_argument("--keep-query", action="store_true", help="Do not strip query strings from URLs (may increase duplicates).")
    p.add_argument("--timeout", type=int, default=DEFAULT_TIMEOUT, help=f"HTTP request timeout in seconds. Default: {DEFAULT_TIMEOUT}")
    p.add_argument("--orientation", choices=["portrait", "landscape"], default="portrait", help="Page orientation for PDF/DOCX/ODT output: portrait (default) or landscape.")
    p.add_argument("--confluence-email", help="Email for Confluence Cloud HTTP Basic Auth.")
    p.add_argument("--confluence-token", help="API Token for Confluence Cloud HTTP Basic Auth.")
    return p.parse_args(argv)


def derive_output_path(start_url: str, fmt: str) -> str:
    netloc = urlparse(start_url).netloc or "output"
    base = re.sub(r"[^A-Za-z0-9_.-]", "_", netloc)
    return f"{base}.{fmt}"


def main(argv: Optional[List[str]] = None) -> int:
    args = parse_args(argv)

    output_format = args.format
    if not output_format:
        if args.output:
            output_format = guess_format_from_extension(args.output)
            if not output_format:
                print("Cannot infer format from output extension. Please specify --format.", file=sys.stderr)
                return 2
        else:
            # default to docx if nothing specified
            output_format = "docx"

    output_path = args.output or derive_output_path(args.url, output_format)

    print(f"[website2docs] Starting crawl: {args.url}")
    print(f"[website2docs] Max pages: {args.max_pages} | Max depth: {args.max_depth} | Delay: {args.delay}s")

    if args.confluence_email and args.confluence_token:
        print("[website2docs] Using Confluence Cloud REST API mode.")
        pages = crawl_confluence(
            start_url=args.url,
            email=args.confluence_email,
            token=args.confluence_token,
            max_depth=args.max_depth,
            max_pages=args.max_pages,
            delay=args.delay,
            timeout=args.timeout,
        )
    else:
        pages = crawl(
            start_url=args.url,
            max_depth=args.max_depth,
            max_pages=args.max_pages,
            delay=args.delay,
            user_agent=args.user_agent,
            keep_query=args.keep_query,
            timeout=args.timeout,
        )

    if not pages:
        print("No pages were fetched. Exiting.", file=sys.stderr)
        return 1

    print(f"[website2docs] Fetched {len(pages)} page(s). Writing {output_format.upper()} -> {output_path}")

    if output_format == "docx":
        save_as_docx(pages, output_path, start_url=args.url, orientation=args.orientation)
    elif output_format == "odt":
        save_as_odt(pages, output_path, start_url=args.url, orientation=args.orientation)
    elif output_format == "pdf":
        save_as_pdf(pages, output_path, start_url=args.url, orientation=args.orientation)
    elif output_format == "html":
        save_as_html(pages, output_path, start_url=args.url, orientation=args.orientation)
    elif output_format == "md":
        save_as_markdown(pages, output_path, start_url=args.url, orientation=args.orientation)
    else:  # pragma: no cover
        print(f"Unsupported format: {output_format}", file=sys.stderr)
        return 2

    print(f"[website2docs] Done: {output_path}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())